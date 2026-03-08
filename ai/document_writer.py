import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import Workbook
from openpyxl.chart import BarChart, PieChart, Reference
from openpyxl.utils import get_column_letter


def save_as_docx(content: str, filepath: str, title: str = ""):
    """Convert markdown-style AI output to a formatted .docx file."""
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    if title:
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    lines = content.split("\n")
    for line in lines:
        stripped = line.strip()

        if not stripped:
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip("*# "), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip("*# "), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip("*# "), level=1)
        elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            doc.add_heading(stripped.strip("*"), level=2)
        # Bullet points
        elif stripped.startswith("* ") or stripped.startswith("- ") or stripped.startswith("• "):
            text = _clean_md(stripped[2:].strip())
            doc.add_paragraph(text, style='List Bullet')
        # Numbered lists
        elif re.match(r'^\d+\.\s', stripped):
            text = _clean_md(re.sub(r'^\d+\.\s*', '', stripped))
            doc.add_paragraph(text, style='List Number')
        # Horizontal rule — skip
        elif stripped.startswith("---"):
            continue
        else:
            text = _clean_md(stripped)
            if text:
                doc.add_paragraph(text)

    doc.save(filepath)


def save_as_text(content: str, filepath: str):
    """Save content as plain .txt file."""
    # Clean markdown formatting for plain text
    clean = content
    clean = re.sub(r'#{1,6}\s*', '', clean)  # Remove heading markers
    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', clean)  # Bold
    clean = re.sub(r'\*(.+?)\*', r'\1', clean)  # Italic
    clean = re.sub(r'`(.+?)`', r'\1', clean)  # Inline code
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(clean)


def save_as_code(content: str, filepath: str):
    """Save code content, stripping any markdown wrappers."""
    code = content.strip()
    # Remove markdown code block wrappers if present
    if code.startswith("```"):
        lines = code.split("\n")
        # Remove first line (```python) and last line (```)
        if lines[-1].strip() == "```":
            lines = lines[1:-1]
        elif lines[0].startswith("```"):
            lines = lines[1:]
        code = "\n".join(lines)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(code)


def save_as_xlsx(content: str, filepath: str, title: str = ""):
    """Convert AI analysis output to .xlsx with charts."""
    wb = Workbook()
    ws = wb.active
    ws.title = title[:31] if title else "Analysis"

    lines = content.split("\n")
    row = 1

    if title:
        ws.cell(row=row, column=1, value=title)
        ws.cell(row=row, column=1).font = ws.cell(row=row, column=1).font.copy(bold=True, size=14)
        row += 2

    # Track tables for chart generation
    tables = []
    current_table_start = None
    current_table_headers = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current_table_start is not None:
                tables.append({
                    "start": current_table_start,
                    "end": row - 1,
                    "headers": current_table_headers
                })
                current_table_start = None
                current_table_headers = []
            row += 1
            continue

        # Markdown table rows
        if "|" in stripped and stripped.count("|") > 1:
            # Clean up leading/trailing pipes if present
            tbl_line = stripped
            if tbl_line.startswith("|"):
                tbl_line = tbl_line[1:]
            if tbl_line.endswith("|"):
                tbl_line = tbl_line[:-1]
                
            cells = [c.strip() for c in tbl_line.split("|")]
            # Skip separator rows
            if all(re.match(r'^[-:]+$', c) for c in cells if c):
                continue
            if current_table_start is None:
                current_table_start = row
                current_table_headers = cells
            for col_idx, cell_val in enumerate(cells, 1):
                clean = _clean_md(cell_val)
                # Try to convert numbers
                try:
                    num = float(clean.replace(",", ""))
                    ws.cell(row=row, column=col_idx, value=num)
                except ValueError:
                    ws.cell(row=row, column=col_idx, value=clean)
                # Bold headers
                if row == current_table_start:
                    ws.cell(row=row, column=col_idx).font = ws.cell(row=row, column=col_idx).font.copy(bold=True)
            row += 1

        # Headings (## Heading)
        elif stripped.startswith("#"):
            if current_table_start is not None:
                tables.append({
                    "start": current_table_start,
                    "end": row - 1,
                    "headers": current_table_headers
                })
                current_table_start = None
                current_table_headers = []
            text = _clean_md(stripped.lstrip("#").strip())
            ws.cell(row=row, column=1, value=text)
            ws.cell(row=row, column=1).font = ws.cell(row=row, column=1).font.copy(bold=True, size=12)
            row += 1
        # Bold text acting as heading (**Heading**)
        elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            if current_table_start is not None:
                tables.append({
                    "start": current_table_start,
                    "end": row - 1,
                    "headers": current_table_headers
                })
                current_table_start = None
                current_table_headers = []
            text = _clean_md(stripped)
            ws.cell(row=row, column=1, value=text)
            ws.cell(row=row, column=1).font = ws.cell(row=row, column=1).font.copy(bold=True, size=11)
            row += 1
        # Bullet points
        elif stripped.startswith("* ") or stripped.startswith("- ") or stripped.startswith("• "):
            text = _clean_md(stripped[2:].strip())
            if ":" in text:
                parts = text.split(":", 1)
                ws.cell(row=row, column=1, value=parts[0].strip())
                ws.cell(row=row, column=2, value=parts[1].strip())
            else:
                ws.cell(row=row, column=1, value=text)
            row += 1
        # Numbered lists
        elif re.match(r'^\d+\.\s', stripped):
            text = _clean_md(re.sub(r'^\d+\.\s*', '', stripped))
            if ":" in text:
                parts = text.split(":", 1)
                ws.cell(row=row, column=1, value=parts[0].strip())
                ws.cell(row=row, column=2, value=parts[1].strip())
            else:
                ws.cell(row=row, column=1, value=text)
            row += 1
        else:
            text = _clean_md(stripped)
            ws.cell(row=row, column=1, value=text)
            row += 1

    # Close any remaining table
    if current_table_start is not None:
        tables.append({
            "start": current_table_start,
            "end": row - 1,
            "headers": current_table_headers
        })

    # Auto-fit column widths
    for col in ws.columns:
        max_length = 0
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        adjusted = min(max_length + 2, 50)
        ws.column_dimensions[col[0].column_letter].width = adjusted

    # Generate charts for tables that have numeric data
    chart_sheet_num = 1
    for table_info in tables:
        _try_add_charts(wb, ws, table_info, chart_sheet_num)
        chart_sheet_num += 1

    wb.save(filepath)


def _try_add_charts(wb, data_ws, table_info, chart_num):
    """Try to create bar and pie charts from table data."""
    start = table_info["start"]
    end = table_info["end"]
    headers = table_info["headers"]

    if end - start < 2 or len(headers) < 2:
        return

    # Find columns with numeric data
    numeric_cols = []
    for col_idx in range(2, len(headers) + 1):
        has_number = False
        for r in range(start + 1, end + 1):
            val = data_ws.cell(row=r, column=col_idx).value
            if isinstance(val, (int, float)):
                has_number = True
                break
        if has_number:
            numeric_cols.append(col_idx)

    if not numeric_cols:
        return

    chart_name = "Charts " + str(chart_num)
    if chart_name in wb.sheetnames:
        chart_ws = wb[chart_name]
    else:
        chart_ws = wb.create_sheet(title=chart_name)

    chart_row = 1

    # Bar chart
    try:
        bar = BarChart()
        bar.type = "col"
        bar.title = headers[0] + " Comparison"
        bar.y_axis.title = "Value"
        bar.x_axis.title = headers[0]

        cats = Reference(data_ws, min_col=1, min_row=start + 1, max_row=end)
        for nc in numeric_cols:
            data = Reference(data_ws, min_col=nc, min_row=start, max_row=end)
            bar.add_data(data, titles_from_data=True)
        bar.set_categories(cats)
        bar.width = 20
        bar.height = 12

        chart_ws.add_chart(bar, "A" + str(chart_row))
        chart_row += 18
    except Exception:
        pass

    # Pie chart (first numeric column only)
    try:
        pie = PieChart()
        pie.title = headers[0] + " Distribution"

        data = Reference(data_ws, min_col=numeric_cols[0], min_row=start, max_row=end)
        cats = Reference(data_ws, min_col=1, min_row=start + 1, max_row=end)
        pie.add_data(data, titles_from_data=True)
        pie.set_categories(cats)
        pie.width = 16
        pie.height = 12

        chart_ws.add_chart(pie, "A" + str(chart_row))
    except Exception:
        pass


def _clean_md(text: str) -> str:
    """Remove markdown formatting from text."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text
